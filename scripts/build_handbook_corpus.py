"""Build the searchable handbook corpus used by the offline Streamlit copilot.

Run this after changing any Word document in ``Bosch_Handbook``.  The source
documents stay in the repository for readers, while the generated JSON keeps
the running application fast and dependency-light.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
HANDBOOK_DIR = PROJECT_ROOT / "Bosch_Handbook"
OUTPUT_PATH = PROJECT_ROOT / "docs" / "handbook" / "bosch_handbook_corpus.json"


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def source_label(path: Path) -> str:
    return f"Bosch Handbook — {path.stem.replace('_', ' ')}"


def document_sections(path: Path) -> list[dict[str, str]]:
    document = Document(path)
    sections: list[dict[str, str]] = []
    heading = "Introduction"
    buffer: list[str] = []

    def flush() -> None:
        text = clean_text(" ".join(buffer))
        if len(text.split()) >= 20:
            sections.append({"heading": heading, "text": text})
        buffer.clear()

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue
        if paragraph.style.name.startswith("Heading"):
            flush()
            heading = text
        else:
            buffer.append(text)
    flush()

    # Important definitions and controls often live in tables, so include them
    # as compact searchable passages as well.
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(" | ".join(cells))
        table_text = clean_text(" ".join(rows))
        if len(table_text.split()) >= 20:
            sections.append({"heading": f"Table {table_index}", "text": table_text})
    return sections


def chunk_text(text: str, words_per_chunk: int = 130) -> list[str]:
    words = text.split()
    return [
        " ".join(words[start : start + words_per_chunk])
        for start in range(0, len(words), words_per_chunk)
        if len(words[start : start + words_per_chunk]) >= 20
    ]


def build_corpus() -> list[dict[str, str]]:
    chunks: list[dict[str, str]] = []
    for path in sorted(HANDBOOK_DIR.glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        for section in document_sections(path):
            for text in chunk_text(section["text"]):
                chunks.append(
                    {
                        "source": source_label(path),
                        "file": path.name,
                        "heading": section["heading"],
                        "text": text,
                    }
                )
    return chunks


def main() -> None:
    if not HANDBOOK_DIR.exists():
        raise FileNotFoundError(f"Handbook folder not found: {HANDBOOK_DIR}")
    corpus = build_corpus()
    if not corpus:
        raise ValueError("No searchable handbook content was found.")
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PATH.write_text(json.dumps(corpus, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {len(corpus)} handbook chunks to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
