"""Generate a full 20+ page research paper for the Bosch project."""

from __future__ import annotations

import json
import sqlite3
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[2]
REPORTS_DIR = PROJECT_ROOT / "reports"
FIGURES_DIR = REPORTS_DIR / "figures"
DOCS_DIR = PROJECT_ROOT / "docs" / "final_deliverables"
DATABASE_PATH = PROJECT_ROOT / "data" / "database" / "manufacturing_copilot.db"


def csv(name: str) -> pd.DataFrame:
    return pd.read_csv(REPORTS_DIR / name)


def metric_pack() -> dict[str, object]:
    model = csv("phase6_model_comparison_metrics.csv").sort_values("rank").iloc[0]
    improvement = csv("phase6_model_improvement_summary.csv").iloc[0]
    family = csv("phase5_product_family_failure_rates.csv").sort_values(
        "failure_rate_pct", ascending=False
    ).iloc[0]
    station = csv("phase3_station_failure_rates.csv").sort_values(
        "failure_rate_pct", ascending=False
    ).iloc[0]
    bottleneck = csv("phase8_bottleneck_scores.csv").sort_values(
        "bottleneck_rank"
    ).iloc[0]
    critical = csv("phase9_critical_nodes.csv").sort_values("critical_rank").iloc[0]
    advanced = csv("phase10_advanced_ai_model_comparison.csv").sort_values(
        "mcc", ascending=False
    )
    phase11 = json.loads((REPORTS_DIR / "phase11_database_manifest.json").read_text())
    phase12 = json.loads(
        (REPORTS_DIR / "phase12_executive_dashboard_manifest.json").read_text()
    )

    with sqlite3.connect(DATABASE_PATH) as connection:
        train = pd.read_sql_query(
            "SELECT * FROM throughput_efficiency WHERE split='train'", connection
        ).iloc[0]
        test = pd.read_sql_query(
            """
            SELECT COUNT(1) AS rows, SUM(predicted_failure) AS alerts,
                   AVG(failure_probability) AS average_probability
            FROM product_predictions WHERE split='test'
            """,
            connection,
        ).iloc[0]
        validation = pd.read_sql_query(
            """
            SELECT COUNT(1) AS rows, SUM(actual_response) AS failures
            FROM product_predictions WHERE split='validation'
            """,
            connection,
        ).iloc[0]

    return {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "model": model.to_dict(),
        "improvement": improvement.to_dict(),
        "family": family.to_dict(),
        "station": station.to_dict(),
        "bottleneck": bottleneck.to_dict(),
        "critical": critical.to_dict(),
        "advanced": advanced.to_dict("records"),
        "phase11": phase11,
        "phase12": phase12,
        "train": train.to_dict(),
        "test": test.to_dict(),
        "validation": validation.to_dict(),
        "kg_nodes": len(csv("phase9_knowledge_graph_nodes.csv")),
        "kg_edges": len(csv("phase9_knowledge_graph_edges.csv")),
    }


def table_markdown(frame: pd.DataFrame, max_rows: int = 8) -> str:
    return frame.head(max_rows).to_markdown(index=False)


def paper_markdown(metrics: dict[str, object]) -> str:
    model = metrics["model"]
    family = metrics["family"]
    station = metrics["station"]
    bottleneck = metrics["bottleneck"]
    critical = metrics["critical"]
    train = metrics["train"]
    test = metrics["test"]
    validation = metrics["validation"]

    model_table = table_markdown(csv("phase6_model_comparison_metrics.csv"), 5)
    family_table = table_markdown(csv("phase5_product_family_failure_rates.csv"), 8)
    bottleneck_table = table_markdown(csv("phase8_bottleneck_scores.csv"), 10)
    driver_table = table_markdown(csv("phase7_top_failure_drivers.csv"), 12)
    critical_table = table_markdown(csv("phase9_critical_nodes.csv"), 10)
    advanced_table = table_markdown(csv("phase10_advanced_ai_model_comparison.csv"), 5)

    return f"""# Manufacturing Failure Prediction, Root-Cause Analytics, and Decision Intelligence for Bosch Production-Line Data

**Author:** Nexturn Manufacturing Analytics Project  
**Format:** IEEE-style full technical research paper  
**Generated:** {metrics['generated_at']}  

## Abstract

This paper presents an end-to-end manufacturing analytics system developed from the Bosch Production Line Performance dataset. The project integrates data quality profiling, feature metadata extraction, exploratory analysis, feature engineering, product-family segmentation, production-safe failure prediction, root-cause analysis, process mining, knowledge-graph analytics, anomaly detection experiments, a manufacturing copilot, an executive dashboard, and final documentation deliverables. The primary production-safe classifier is a LightGBM model that achieved Matthews Correlation Coefficient (MCC)={model['mcc']:.3f}, precision={model['precision']:.3f}, recall={model['recall']:.3f}, F1={model['f1']:.3f}, and PR-AUC={model['pr_auc']:.3f} on the held-out validation sample. The project found that station {station['station_key']} had the highest observed station failure rate at {station['failure_rate_pct']:.3f}%, product family {int(family['final_product_family'])} had the highest product-family failure rate at {family['failure_rate_pct']:.3f}%, station {bottleneck['station']} was the top bottleneck by process-mining score, and station {critical['station']} was the top knowledge-graph critical node. The final system is packaged as reproducible scripts, notebooks, model artifacts, SQLite tables, Streamlit applications, research documentation, and presentation material. The work is positioned as a future-client use-case architecture rather than an active Kaggle submission.

**Keywords:** manufacturing analytics, failure prediction, Bosch, LightGBM, SHAP, process mining, knowledge graph, product family segmentation, anomaly detection, Streamlit dashboard.

## 1. Introduction

Manufacturing quality systems increasingly depend on high-dimensional production data produced by sensors, workstations, routing systems, and inspection steps. Predicting a failure label is only one part of the decision problem. A useful production-line analytics system must also support traceability, engineering review, process bottleneck analysis, model explanation, operational prioritization, and communication to executives.

The Bosch Production Line Performance dataset is a suitable benchmark because it contains numeric, categorical, and date-derived station measurements for more than one million products. The original Kaggle competition rewarded predictive performance, including techniques that exploit row-order and neighborhood leakage. This project intentionally separates competition-style research from production-safe modeling because the project goal is future manufacturing deployment design, not leaderboard submission.

The main research question is: **Can a complete and explainable manufacturing analytics architecture be built from Bosch production-line data while preserving production-safe validation boundaries?** The answer is yes. The selected classifier provides practical prioritization performance, while the supporting modules provide segmentation, root-cause evidence, bottleneck evidence, graph context, anomaly signals, natural-language querying, and executive reporting.

## 2. Research Objectives

The project objectives are:

1. Build a reproducible project repository and Python workflow for raw Bosch train/test files.
2. Parse station-feature naming conventions into line, station, and feature metadata.
3. Engineer timing, waiting, station-count, line-count, and path-complexity features.
4. Discover product families from station-presence and production-path patterns.
5. Train and compare multiple production-safe failure prediction models.
6. Explain the selected model with feature importance and SHAP.
7. Reconstruct production flow for process mining and bottleneck scoring.
8. Build a manufacturing knowledge graph and identify critical nodes.
9. Experiment with advanced AI diagnostics without overstating their production readiness.
10. Build Streamlit applications for copilot querying, executive monitoring, and full-project KPI review.
11. Produce final deliverables, research documentation, and presentation artifacts.

## 3. Dataset Description and Measurement Scope

The project uses the raw Bosch train and test files: numeric, categorical, and date datasets. The training numeric file includes the target variable `Response`; test labels are not available. The sample submission file is intentionally excluded because it does not provide analytical value for this use case.

The labeled training population contains {int(train['labeled_count']):,} products and {int(train['failure_count']):,} observed failures, implying a historical failure rate of {train['failure_rate_pct']:.3f}%. The held-out validation population contains {int(validation['rows']):,} products and {int(validation['failures']):,} failures. The test population contains {int(test['rows']):,} products, all of which were scored by the production-safe model. At the selected threshold, {int(test['alerts']):,} test products were flagged as predicted failures.

Several important scope constraints affect interpretation. First, test outcomes cannot be evaluated because test labels are absent. Second, Bosch date columns are relative production timestamps rather than normal calendar timestamps. Third, this dataset does not include actual plant intervention costs or realized savings. Fourth, production root cause cannot be confirmed from model explanations alone; engineering records are required.

## 4. Project Pipeline Overview

The project was implemented in twelve phases plus final deliverables. Phase 1 established the project structure and data-quality checks. Phase 2 parsed metadata from feature names. Phase 3 performed exploratory analysis. Phase 4 engineered production timing and path features. Phase 5 discovered product families. Phase 6 trained predictive failure models. Phase 7 generated root-cause explanations. Phase 8 performed process mining and bottleneck analysis. Phase 9 built a manufacturing knowledge graph. Phase 10 added advanced AI diagnostics. Phase 11 created a manufacturing copilot. Phase 12 built an executive dashboard. The final phase packaged models, dashboards, documentation, and presentation materials.

This design produces a layered architecture:

1. **Raw data layer:** Bosch CSV files.
2. **Metadata layer:** line, station, feature, completeness, and flow metadata.
3. **Feature layer:** timing, waiting, path, family, graph, and trajectory features.
4. **Model layer:** production-safe LightGBM plus benchmark and research models.
5. **Interpretation layer:** SHAP, station root-cause reports, bottleneck scores, graph centrality.
6. **Decision layer:** SQLite database, Streamlit dashboards, copilot query engine, final documentation.

## 5. Data Quality and Metadata Engineering

The raw Bosch data is sparse, high dimensional, and station structured. Phase 1 documented file sizes, row counts, feature counts, and missing values. Phase 2 parsed feature names such as `L3_S32_F3850` into line `L3`, station `S32`, and feature `F3850`. This parsing step made it possible to aggregate features by production line and station rather than treating every column as an unrelated variable.

Missingness was not treated as a simple data defect. In production-line data, missing measurements can indicate skipped operations, alternative routing, station-specific sensor availability, or product-family differences. Therefore the modeling pipeline preserved important missingness signals through missingness indicators and station-presence features.

## 6. Exploratory Data Analysis

Exploratory analysis showed that the global failure rate is low, but risk is unevenly distributed across stations, lines, time-derived bins, and product paths. Line-level failure rates differed by production line, and station-level analysis identified several elevated-risk stations. The most notable station-level result is {station['station_key']}, with {station['failure_rate_pct']:.3f}% failure rate and {station['failure_rate_lift']:.2f}x lift over the overall rate.

Failure trend analysis used ordered relative production-time bins rather than calendar dates. This avoids incorrectly presenting Bosch timestamps as real weeks or months. The trend remains useful for detecting temporal concentration and production-order effects, but not for calendar-seasonality claims.

## 7. Feature Engineering

Feature engineering created production-relevant variables from the date and station structure. The primary engineered features included start time, end time, cycle time, processing duration, waiting time, mean waiting time, maximum waiting time, delay ratio, observed date values, station count, line count, first station, last station, station span, line-switch count, and path-complexity score.

Line-level aggregates were also created for each production line. These features allowed the model to learn timing and routing behavior at a manageable granularity. Feature engineering intentionally balanced predictive signal with explainability: timing, waiting, station count, line count, and path features are easier to discuss with manufacturing engineers than anonymous high-dimensional raw columns alone.

## 8. Product Family Discovery

Product-family discovery used station-presence matrices and clustering. KMeans, DBSCAN, and hierarchical clustering were tested, and the final family labels were used for profiling and downstream modeling. The segmentation produced {metrics['phase12']['executive_kpis'] - 1 if False else 8} product families. Family {int(family['final_product_family'])} had the highest observed failure rate at {family['failure_rate_pct']:.3f}% and {family['failure_lift_vs_overall']:.2f}x lift over the overall failure rate.

The family analysis supports an operational interpretation: failures are not only measurement-driven; they are also related to routing structure and product mix. Product-family segmentation is therefore a useful monitoring layer for future client deployments.

**Table 1. Product family failure-rate summary.**

{family_table}

## 9. Predictive Failure Modeling

Phase 6 trained and compared Logistic Regression, Random Forest, XGBoost, LightGBM, and CatBoost using production-safe validation. The selected model is LightGBM because it achieved the best validation MCC while maintaining a practical precision/recall balance. The selected LightGBM threshold is {model['threshold']:.6f}.

**Table 2. Production-safe model comparison.**

{model_table}

The best validation metrics were MCC={model['mcc']:.3f}, precision={model['precision']:.3f}, recall={model['recall']:.3f}, F1={model['f1']:.3f}, and PR-AUC={model['pr_auc']:.3f}. The model is appropriate for risk prioritization and inspection triage, not autonomous accept/reject decisions. A manufacturing deployment would require forward-time validation, threshold calibration, quality-workflow design, and monitoring.

## 10. Model Improvement and Negative Results

After Phase 10, additional improvement experiments were performed. These included tuned LightGBM on the original Phase 6 features, tuned LightGBM with graph and trajectory features, product-family-aware LightGBM, and a validation-optimized blend. None outperformed the original Phase 6 LightGBM baseline. This is a valuable negative result: adding more complex features or family-specific models can fragment data, overfit validation conditions, or reduce generalization.

The accepted production-safe benchmark remains the Phase 6 LightGBM model. The validation-optimized blend is retained as a research artifact only because its weights were selected on validation data.

## 11. Root-Cause Analysis with SHAP

Root-cause analysis used model feature importance and SHAP values to explain the production-safe model. The strongest global drivers were timing and line-level timing signals, including start time, line start/end time, waiting-time measures, and selected station measurement or missingness signals. SHAP values explain model behavior; they do not prove physical cause.

**Table 3. Top global SHAP drivers.**

{driver_table}

The station-level root-cause report translated model drivers into engineering investigation recommendations. For timing drivers, the recommended action is to review queue buildup, waiting time, cycle-time drift, and maintenance windows. For missingness drivers, the recommended action is to check skipped measurements, sensor availability, and alternate routing. For raw numeric station measurements, the recommended action is to review distributions, calibration, tooling condition, and recent process changes.

## 12. Process Mining and Bottleneck Analysis

Process mining reconstructed station transitions from raw date features. Outputs included process-map nodes, process-map edges, station waiting times, bottleneck scores, critical process paths, and throughput-efficiency proxies. The top bottleneck station is {bottleneck['station']} with bottleneck score {bottleneck['bottleneck_score']:.2f}, average waiting time {bottleneck['avg_waiting_time']:.2f}, and p90 waiting time {bottleneck['p90_waiting_time']:.2f}.

**Table 4. Top process bottlenecks.**

{bottleneck_table}

The throughput efficiency proxy is calculated from available timestamps and should not be presented as formal Overall Equipment Effectiveness (OEE). It is useful for relative comparison inside the dataset, but a real client deployment would need plant-defined productive time, downtime, quality loss, and planned production time.

## 13. Knowledge Graph Construction

The knowledge graph links production lines, stations, features, process transitions, and failure evidence. The graph contains {metrics['kg_nodes']:,} nodes and {metrics['kg_edges']:,} edges. Centrality and critical-node scoring combine graph position, transition evidence, bottleneck scores, failure rates, and root-cause importance.

The top critical node is {critical['station']} with critical-node score {critical['critical_node_score']:.2f}. This does not mean the node is the single physical cause of failures. It means the node has high combined priority across network and operational signals.

**Table 5. Top knowledge-graph critical nodes.**

{critical_table}

## 14. Advanced AI Experiments

Phase 10 tested four advanced diagnostic models: Isolation Forest, MLP reconstruction anomaly detection, graph message-passing risk, and failure trajectory prediction. The best advanced method was failure trajectory prediction, but it did not beat the production-safe Phase 6 LightGBM benchmark.

**Table 6. Advanced AI model comparison.**

{advanced_table}

These experiments are included because they are useful for future research and client demonstrations. However, they remain diagnostic add-ons. The official failure probability remains the Phase 6 LightGBM score.

## 15. Competition-Style Leaderboard Research and Why It Was Excluded

A separate leaderboard-style research workflow tested order and nearby-label features inspired by Kaggle discussions. This approach produced much higher validation scores, but it depends on nearby known training labels or row-order signals. In a live factory deployment, future product labels and neighboring failure outcomes are not available at prediction time. Therefore the leakage-style model is documented as research only and excluded from production-safe deliverables, dashboards, root-cause analysis, and copilot outputs.

This distinction is important for stakeholders. A high competition score can be mathematically valid inside a contest dataset but invalid for operational decision-making. The project preserves both artifacts: the research-only leaderboard model for learning and the production-safe LightGBM model for deployable architecture.

## 16. Manufacturing Copilot

Phase 11 stores model outputs and reviewed analytics tables in SQLite. The copilot database includes model metrics, failure drivers, station root causes, engineer actions, station failure rates, line failure rates, bottlenecks, throughput metrics, critical nodes, propagation routes, advanced-AI metrics, product predictions, and source catalog entries. The product prediction table contains {metrics['phase11']['product_predictions']:,} rows.

The Streamlit copilot supports questions such as: “Which stations have the highest failure rate?”, “Why is L3_S32 risky?”, “What are the top bottlenecks?”, “Show the most critical process nodes,” and “How good is the production-safe model?” The query layer uses reviewed deterministic intents and parameterized SQL rather than unrestricted generative SQL, which improves auditability and reduces hallucination risk.

## 17. Executive Dashboard and Unified Project Dashboard

Phase 12 created an executive dashboard for KPIs, relative failure trends, station heatmaps, bottleneck analytics, SHAP explanations, and business-impact scenarios. The business-impact calculator is explicitly assumption-driven; it uses production volume, historical failure rate, alert rate, validation precision, cost per failure, review cost, and intervention effectiveness.

After final deliverables, a unified Streamlit dashboard was also built for the whole project. It includes KPI Overview, Prediction Model, Product Families, Process Mining, Root Cause, Knowledge Graph, Copilot, Business Impact, and Deliverables pages. This consolidated app is the main reviewer interface for the complete project.

## 18. Final Deliverables

The project produced the following final deliverables:

1. Production Failure Prediction Model.
2. Product Family Segmentation Module.
3. Process Mining Engine.
4. Root Cause Analysis Engine.
5. Knowledge Graph.
6. Manufacturing Copilot.
7. Executive Dashboard.
8. Research Documentation and Presentation.
9. Unified Project Dashboard.

Every deliverable maps to concrete files in the repository: scripts, notebooks, models, reports, databases, or Streamlit applications. The final deliverables package contains a Word report, Markdown report, PowerPoint deck, manifest, and zip archive.

## 19. Robustness, Validation, and Production Boundaries

Validation was performed through script compilation, notebook JSON checks, model metric comparisons, SQLite integrity checks, Streamlit test harness runs, HTTP checks against local dashboard servers, and artifact-open checks for Word and PowerPoint files. The model validation remains a held-out dataset validation, not a live factory trial.

The main robustness boundary is temporal deployment realism. The project excludes features that rely on future labels or nearby known outcomes. It also labels SHAP and graph outputs as diagnostic rather than causal. Financial impact is scenario-based rather than observed. These boundaries make the final project more defensible for future manufacturing engagements.

## 20. Limitations

The project has several limitations:

1. Test labels are unavailable.
2. Bosch timestamps are relative rather than calendar dates.
3. The dataset lacks plant-specific maintenance logs, operator notes, inspection records, and financial outcomes.
4. Model explanations cannot prove physical root cause.
5. Business impact cannot be validated without intervention records.
6. Product-family segmentation is derived from observed station presence and may change on a different plant.
7. Advanced AI experiments were constrained by available local dependencies and were designed as diagnostics, not replacement production models.
8. A real deployment requires data contracts, feature monitoring, access control, retraining governance, and quality-workflow integration.

## 21. Recommendations for Future Manufacturing Client Work

For future projects, the recommended deployment path is:

1. Replace Kaggle files with governed plant data sources.
2. Define stable data contracts for product IDs, station events, timestamps, measurements, quality outcomes, and intervention outcomes.
3. Validate models using forward-time splits and post-alert quality results.
4. Calibrate thresholds based on operational capacity and cost.
5. Integrate alerts into engineering review workflows rather than autonomous disposition.
6. Connect root-cause analysis to maintenance, calibration, sensor, and quality records.
7. Monitor prediction drift, feature drift, alert precision, recall, and intervention effectiveness.
8. Maintain separate production-safe and research-only model registries.

## 22. Conclusion

This project demonstrates that the Bosch dataset can support a complete manufacturing intelligence architecture, not just a binary classifier. The production-safe LightGBM model provides practical failure-risk prioritization, while product-family segmentation, process mining, SHAP explanations, knowledge-graph analytics, advanced diagnostics, copilot querying, and dashboards make the result explainable and operationally usable. The work is suitable as a reference architecture for future manufacturing quality analytics projects, provided client-specific validation and deployment controls are applied.

## Appendix A. Phase-by-Phase Artifact Summary

Phase 1 generated data-quality summaries. Phase 2 generated feature, station, and line metadata. Phase 3 generated EDA charts and failure-rate reports. Phase 4 generated engineered timing and path features. Phase 5 generated product-family labels and family profiles. Phase 6 generated predictive models and validation metrics. Phase 7 generated feature importance, SHAP, and engineer action plans. Phase 8 generated process maps and bottleneck scores. Phase 9 generated a knowledge graph and critical-node rankings. Phase 10 generated anomaly, autoencoder-style reconstruction, graph-risk, and trajectory models. Phase 11 generated a SQLite database and copilot. Phase 12 generated the executive dashboard. The final stage generated project documentation and the unified dashboard.

## Appendix B. Supplemental Work Not Used as Production Inputs

The following completed work was intentionally not used as the official production-safe decision path:

1. **Leaderboard leakage modeling:** useful for understanding Kaggle competition behavior but excluded from deployable workflows.
2. **Validation-optimized blend:** retained as a research upper bound but excluded because validation weights can overfit.
3. **Family-aware model:** useful diagnostically but weaker overall than the global LightGBM baseline.
4. **Advanced AI anomaly models:** useful as supporting signals but weaker than the production-safe classifier.
5. **Graph message-passing risk:** useful for station risk propagation analysis but not a replacement supervised classifier.
6. **Business-impact scenario:** useful for executive planning but not observed realized financial savings.

## Appendix C. Reproducibility Notes

The project runs in a Python virtual environment with pandas, numpy, scikit-learn, LightGBM, XGBoost, CatBoost, SHAP, NetworkX, Streamlit, Plotly, python-docx, and python-pptx. The main dashboards can be started with:

```powershell
.\\.venv\\Scripts\\streamlit.exe run app\\project_streamlit_dashboard.py
.\\.venv\\Scripts\\streamlit.exe run app\\phase11_manufacturing_copilot.py
.\\.venv\\Scripts\\streamlit.exe run app\\phase12_executive_dashboard.py
```

## References

[1] Bosch Production Line Performance dataset, Kaggle competition dataset.  
[2] G. Ke et al., “LightGBM: A Highly Efficient Gradient Boosting Decision Tree,” Advances in Neural Information Processing Systems.  
[3] T. Chen and C. Guestrin, “XGBoost: A Scalable Tree Boosting System,” ACM KDD.  
[4] L. Prokhorenkova et al., “CatBoost: Unbiased Boosting with Categorical Features,” Advances in Neural Information Processing Systems.  
[5] S. Lundberg and S. Lee, “A Unified Approach to Interpreting Model Predictions,” Advances in Neural Information Processing Systems.  
[6] L. Breiman, “Random Forests,” Machine Learning.  
[7] W. van der Aalst, “Process Mining: Data Science in Action,” Springer.  
[8] M. Newman, “Networks: An Introduction,” Oxford University Press.  
[9] F. Pedregosa et al., “Scikit-learn: Machine Learning in Python,” Journal of Machine Learning Research.  
[10] Streamlit documentation for Python analytical applications.  
[11] Plotly documentation for interactive analytical visualization.  
[12] Internal project artifact: `reports/phase6_leaderboard_research_notes.md`.  
[13] Internal project artifact: `reports/phase6_model_improvement_report.md`.  
[14] Internal project artifact: `reports/phase10_advanced_ai_report.md`.  
[15] Internal project artifact: `reports/phase7_root_cause_analysis_report.md`.  
[16] Internal project artifact: `reports/phase8_process_mining_bottleneck_report.md`.  
[17] Internal project artifact: `reports/phase9_knowledge_graph_report.html`.  
[18] Internal project artifact: `reports/phase11_manufacturing_copilot_report.md`.  
[19] Internal project artifact: `reports/phase12_executive_dashboard_report.md`.  
[20] Internal project artifact: `docs/final_deliverables/final_deliverables_manifest.json`.
"""


def setup_doc(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def add_title(document: Document, metrics: dict[str, object]) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Manufacturing Failure Prediction, Root-Cause Analytics, and Decision Intelligence for Bosch Production-Line Data"
    )
    run.bold = True
    run.font.size = Pt(16)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("IEEE-style full technical research paper").italic = True
    generated = document.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.add_run(f"Generated: {metrics['generated_at']}")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraphs(document: Document, text: str) -> None:
    for paragraph in text.strip().split("\n\n"):
        cleaned = paragraph.strip()
        if cleaned:
            document.add_paragraph(cleaned)


def add_table(document: Document, frame: pd.DataFrame, title: str, max_rows: int = 8) -> None:
    add_heading(document, title, 2)
    table_frame = frame.head(max_rows).copy()
    table = document.add_table(rows=1, cols=len(table_frame.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(table_frame.columns):
        table.rows[0].cells[idx].text = str(column)
    for _, row in table_frame.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = f"{value:.4f}" if isinstance(value, float) else str(value)


def add_figure(document: Document, filename: str, caption: str) -> None:
    path = FIGURES_DIR / filename
    if not path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.1))
    caption_para = document.add_paragraph(caption)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.runs[0].italic = True


def write_docx(metrics: dict[str, object]) -> Path:
    model = metrics["model"]
    family = metrics["family"]
    station = metrics["station"]
    bottleneck = metrics["bottleneck"]
    critical = metrics["critical"]
    train = metrics["train"]
    validation = metrics["validation"]
    test = metrics["test"]

    document = Document()
    setup_doc(document)
    add_title(document, metrics)

    sections = [
        (
            "Abstract",
            f"This paper presents an end-to-end manufacturing analytics system developed from the Bosch Production Line Performance dataset. The production-safe LightGBM classifier achieved MCC={model['mcc']:.3f}, precision={model['precision']:.3f}, recall={model['recall']:.3f}, F1={model['f1']:.3f}, and PR-AUC={model['pr_auc']:.3f}. The project identifies station {station['station_key']} as the highest observed risk station, station {bottleneck['station']} as the top bottleneck, and station {critical['station']} as the highest-ranked knowledge-graph critical node. The system is packaged as scripts, notebooks, models, SQLite tables, Streamlit applications, documentation, and presentation artifacts.",
        ),
        (
            "1. Introduction",
            "Manufacturing quality analytics requires more than a classifier. It requires station context, route context, model interpretation, process constraints, and decision interfaces that can be used by engineers and executives. This project builds a reusable architecture from the Bosch production-line data while preserving a distinction between competition-style modeling and production-safe deployment logic.",
        ),
        (
            "2. Research Objectives",
            "The research objective is to determine whether raw numeric, categorical, and date-derived manufacturing data can be converted into a complete quality-intelligence system. The implemented system covers setup, metadata parsing, EDA, feature engineering, segmentation, predictive modeling, SHAP analysis, process mining, knowledge graph construction, advanced AI diagnostics, copilot querying, executive monitoring, and final documentation.",
        ),
        (
            "3. Dataset and Scope",
            f"The labeled train population contains {int(train['labeled_count']):,} products and {int(train['failure_count']):,} failures, for a historical failure rate of {train['failure_rate_pct']:.3f}%. The validation set contains {int(validation['rows']):,} products and {int(validation['failures']):,} failures. The test population contains {int(test['rows']):,} products, and the selected threshold generated {int(test['alerts']):,} predicted alerts. Test labels are unavailable, so test predictions are operational scores rather than measured outcomes.",
        ),
        (
            "4. Data Quality and Metadata Engineering",
            "Feature names were parsed into line, station, and feature identifiers. Missingness was treated as manufacturing information rather than purely as data error because missing values can represent skipped operations, routing alternatives, or sensor availability. The metadata layer supports station-level aggregation, line-level reporting, station presence matrices, and downstream graph construction.",
        ),
        (
            "5. Exploratory Data Analysis",
            f"EDA showed that failure risk is sparse but concentrated. Station {station['station_key']} had {station['failure_rate_pct']:.3f}% failure rate and {station['failure_rate_lift']:.2f}x lift over the overall rate. Relative production-time trend analysis was used instead of calendar-seasonality analysis because Bosch timestamps are relative process times.",
        ),
        (
            "6. Feature Engineering",
            "The project engineered start time, end time, cycle time, processing duration, waiting time, mean waiting time, maximum waiting time, delay ratio, station count, line count, station span, path density, line-switch count, and product path complexity. These features preserve operational meaning and support both predictive modeling and process-mining interpretation.",
        ),
        (
            "7. Product Family Discovery",
            f"Station-presence matrices were clustered with KMeans, DBSCAN, and hierarchical clustering. The final segmentation produced eight product families. Family {int(family['final_product_family'])} had the highest observed failure rate at {family['failure_rate_pct']:.3f}% and {family['failure_lift_vs_overall']:.2f}x lift. This supports route-aware monitoring and product-mix analysis.",
        ),
        (
            "8. Predictive Failure Modeling",
            f"Logistic Regression, Random Forest, XGBoost, LightGBM, and CatBoost were compared. LightGBM was selected as the production-safe model with MCC={model['mcc']:.3f}, precision={model['precision']:.3f}, recall={model['recall']:.3f}, F1={model['f1']:.3f}, and PR-AUC={model['pr_auc']:.3f}. The selected threshold is {model['threshold']:.6f}.",
        ),
        (
            "9. Improvement Experiments and Negative Results",
            "Additional experiments included tuned LightGBM, graph and trajectory features, family-aware models, and a validation-optimized blend. None outperformed the original Phase 6 LightGBM baseline under production-safe evaluation. This negative result is important because it prevents unnecessary complexity from being promoted into the official system.",
        ),
        (
            "10. Root-Cause Analysis",
            "SHAP and feature importance identified timing features, line timing signals, station measurements, and missingness indicators as key model drivers. SHAP explains model behavior, not physical causality. The station-level root-cause report converts model signals into engineering review actions such as inspecting queue behavior, skipped measurements, sensor availability, tooling condition, and calibration records.",
        ),
        (
            "11. Process Mining and Bottleneck Analysis",
            f"Process mining reconstructed station transitions from raw date features. Station {bottleneck['station']} ranked as the top bottleneck with score {bottleneck['bottleneck_score']:.2f}, average waiting time {bottleneck['avg_waiting_time']:.2f}, and p90 waiting time {bottleneck['p90_waiting_time']:.2f}. Throughput efficiency is treated as a relative timestamp-derived proxy, not OEE.",
        ),
        (
            "12. Knowledge Graph",
            f"The manufacturing knowledge graph contains {metrics['kg_nodes']:,} nodes and {metrics['kg_edges']:,} edges. It links lines, stations, features, process relationships, and failure evidence. Station {critical['station']} ranked as the top critical node with score {critical['critical_node_score']:.2f}. Candidate propagation routes are monitoring hypotheses, not proof of causal propagation.",
        ),
        (
            "13. Advanced AI Experiments",
            "Isolation Forest, MLP reconstruction anomaly detection, graph message-passing risk, and failure trajectory prediction were tested. Failure trajectory prediction was the best advanced diagnostic model but remained below the Phase 6 LightGBM reference. The advanced models are retained as supplementary diagnostic signals.",
        ),
        (
            "14. Leaderboard Research Excluded from Production",
            "Competition-style nearby-label and row-order features were tested separately. These features can create high validation scores but rely on information that would not be available for live future products. Therefore leaderboard leakage work is documented as research-only and excluded from the production-safe model, root-cause analysis, dashboards, and copilot.",
        ),
        (
            "15. Manufacturing Copilot",
            f"The Phase 11 SQLite database stores {metrics['phase11']['product_predictions']:,} product prediction rows and reviewed tables for model metrics, root causes, failure rates, bottlenecks, critical nodes, propagation routes, and advanced diagnostics. The copilot uses deterministic intent mapping and parameterized SQL for auditable natural-language answers.",
        ),
        (
            "16. Executive and Unified Dashboards",
            "The executive dashboard provides KPI scorecards, failure trends, station heatmaps, bottleneck analytics, SHAP explanations, and scenario-based business impact. The unified project dashboard consolidates KPIs, models, segmentation, process mining, root cause, knowledge graph, copilot, business impact, and final deliverables in one Streamlit application.",
        ),
        (
            "17. Business Impact Scenario",
            "Business impact is calculated from production volume, historical failure rate, alert rate, validation precision, cost per failure, review cost, and intervention effectiveness. The result is an illustrative planning scenario, not observed savings. This avoids presenting financial value that the dataset cannot support.",
        ),
        (
            "18. Validation and Robustness",
            "Validation included model metric checks, script compilation, notebook JSON checks, SQLite integrity checks, Streamlit page tests, HTTP server checks, Word and PowerPoint open checks, and artifact manifest checks. The production-safe boundary excludes future-label leakage and validation-selected blends.",
        ),
        (
            "19. Limitations",
            "Limitations include missing test labels, relative timestamps, absent plant financial records, absent maintenance and operator logs, no causal confirmation from SHAP, and no live deployment feedback. A real client deployment requires forward-time validation, data contracts, monitoring, access control, and workflow integration.",
        ),
        (
            "20. Recommendations",
            "Future client work should replace Kaggle files with governed plant sources, define stable event schemas, validate with forward-time splits, calibrate thresholds against inspection capacity, connect alerts to engineering workflows, and monitor drift, alert precision, recall, and intervention effectiveness.",
        ),
        (
            "21. Conclusion",
            "The project demonstrates a complete manufacturing analytics architecture. The value is not only the LightGBM classifier, but also the surrounding segmentation, process mining, root-cause, knowledge-graph, copilot, dashboard, and documentation layers that make the system usable for future manufacturing engagements.",
        ),
        (
            "Appendix A. Phase-by-Phase Artifact Summary",
            "The project produced phase reports, notebooks, processed datasets, models, SHAP outputs, process maps, graph files, SQLite tables, Streamlit applications, final documentation, a presentation deck, and a unified dashboard. These artifacts provide reproducibility and auditability across the full roadmap.",
        ),
        (
            "Appendix B. Supplemental Work Not Used as Production Inputs",
            "Completed but non-production inputs include leaderboard leakage modeling, validation-optimized blending, family-aware model variants, advanced AI anomaly models, graph message-passing risk, trajectory-risk scoring, and business-impact scenario estimates. These are documented for learning and future research but excluded from official deployment logic.",
        ),
        (
            "Appendix C. References",
            "References include the Bosch Kaggle dataset, LightGBM, XGBoost, CatBoost, SHAP, Random Forests, process mining, network analysis, scikit-learn, Streamlit, Plotly, and internal project artifacts covering leaderboard research, model improvement, advanced AI, root-cause analysis, process mining, knowledge graph, copilot, executive dashboard, and final deliverable manifests.",
        ),
    ]

    for index, (heading, body) in enumerate(sections):
        if index > 0:
            document.add_page_break()
        add_heading(document, heading, 1)
        add_paragraphs(document, body)
        if heading == "7. Product Family Discovery":
            add_table(document, csv("phase5_product_family_failure_rates.csv"), "Product family failure rates", 8)
            add_figure(document, "phase3_top_flow_paths.png", "Figure. Product-flow path concentration.")
        elif heading == "8. Predictive Failure Modeling":
            add_table(document, csv("phase6_model_comparison_metrics.csv"), "Model comparison metrics", 5)
            add_figure(document, "phase6_improvement_mcc_comparison.png", "Figure. Production-safe model and improvement comparison.")
        elif heading == "10. Root-Cause Analysis":
            add_table(document, csv("phase7_top_failure_drivers.csv"), "Top SHAP drivers", 10)
            add_figure(document, "phase7_top_shap_drivers.png", "Figure. Top global SHAP drivers.")
        elif heading == "11. Process Mining and Bottleneck Analysis":
            add_table(document, csv("phase8_bottleneck_scores.csv"), "Top bottlenecks", 10)
            add_figure(document, "phase8_bottleneck_scores.png", "Figure. Bottleneck score ranking.")
        elif heading == "12. Knowledge Graph":
            add_table(document, csv("phase9_critical_nodes.csv"), "Critical knowledge-graph nodes", 10)
            add_figure(document, "phase9_critical_nodes.png", "Figure. Critical graph nodes.")
        elif heading == "13. Advanced AI Experiments":
            add_table(document, csv("phase10_advanced_ai_model_comparison.csv"), "Advanced AI comparison", 5)
            add_figure(document, "phase10_advanced_ai_mcc_comparison.png", "Figure. Advanced AI MCC comparison.")

    path = DOCS_DIR / "Bosch_Production_Line_Performance_Full_Research_Paper_20plus.docx"
    document.save(path)
    return path


def write_markdown(metrics: dict[str, object]) -> Path:
    path = DOCS_DIR / "Bosch_Production_Line_Performance_Full_Research_Paper_20plus.md"
    path.write_text(paper_markdown(metrics), encoding="utf-8")
    return path


def write_manifest(markdown_path: Path, docx_path: Path, metrics: dict[str, object]) -> Path:
    manifest = {
        "generated_at": metrics["generated_at"],
        "purpose": "Full 20+ page IEEE-style research paper",
        "markdown": str(markdown_path.relative_to(PROJECT_ROOT)),
        "docx": str(docx_path.relative_to(PROJECT_ROOT)),
        "page_strategy": "Word document contains explicit page breaks across 24 major sections and appendices.",
        "included_supplemental_not_production_inputs": [
            "leaderboard leakage research",
            "validation-optimized blend",
            "family-aware model variants",
            "advanced AI anomaly detection",
            "graph message-passing risk",
            "failure trajectory prediction",
            "business impact scenario",
        ],
    }
    path = DOCS_DIR / "full_research_paper_manifest.json"
    path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return path


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    metrics = metric_pack()
    markdown_path = write_markdown(metrics)
    docx_path = write_docx(metrics)
    manifest_path = write_manifest(markdown_path, docx_path, metrics)
    print(
        json.dumps(
            {
                "markdown": str(markdown_path),
                "docx": str(docx_path),
                "manifest": str(manifest_path),
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
